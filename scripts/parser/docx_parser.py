"""DOCX parser implementation."""
from pathlib import Path
from typing import List
from docx import Document
from scripts.parser.base import BaseParser
from scripts.models import ParsedDocument, Heading, Table, ListItem


class DocxParser(BaseParser):
    """Parser for Word DOCX documents."""
    
    def parse(self, source: str) -> ParsedDocument:
        """Parse a DOCX file into a ParsedDocument.
        
        Args:
            source: Path to the DOCX file.
            
        Returns:
            ParsedDocument with extracted content.
            
        Raises:
            FileNotFoundError: If the DOCX file doesn't exist.
        """
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {source}")
        
        doc = self._create_empty_document()
        doc.raw_text = ""
        
        docx_doc = Document(str(path))
        
        # Process paragraphs
        for para in docx_doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if it's a heading based on style
            style_name = para.style.name if para.style else ""
            
            if style_name.lower().startswith('heading'):
                # Extract heading level
                level_match = style_name.lower()
                if 'heading 1' in level_match:
                    level = 1
                elif 'heading 2' in level_match:
                    level = 2
                elif 'heading 3' in level_match:
                    level = 3
                elif 'heading 4' in level_match:
                    level = 4
                elif 'heading 5' in level_match:
                    level = 5
                elif 'heading 6' in level_match:
                    level = 6
                else:
                    level = 1
                
                doc.headings.append(Heading(level=level, text=text))
                doc.raw_text += text + "\n"
            else:
                # Regular paragraph
                doc.paragraphs.append(text)
                doc.raw_text += text + "\n"
        
        # Process tables
        for table in docx_doc.tables:
            table_obj = self._parse_docx_table(table)
            doc.tables.append(table_obj)
        
        return doc
    
    def _parse_docx_table(self, table) -> Table:
        """Parse a DOCX table element.
        
        Args:
            table: python-docx Table object.
            
        Returns:
            Table with headers and rows.
        """
        headers = []
        rows = []
        
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            
            if i == 0:
                headers = cells
            else:
                rows.append(cells)
        
        return Table(headers=headers, rows=rows)
